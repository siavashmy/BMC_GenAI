import streamlit as st
import google.generativeai as genai
import os
from visual_business_model_canvas import show_bmc_visualization
from io import BytesIO
from docx import Document

# -------------------------------
# Configure Gemini API
# -------------------------------
if "GEMINI_API_KEY" not in st.secrets:
    st.error("❌ Please add your Gemini API key in Streamlit Secrets.")
    st.stop()

genai.configure(api_key=st.secrets["GEMINI_API_KEY"])
model = genai.GenerativeModel("gemini-2.0-flash")

# -------------------------------
# Define workflow steps
# -------------------------------
STEPS = [
    "Story Input",
    "Focus Generation",
    "Issues Generation",
    "Tension Matrix",
    "Dilemmas & Ranking",
    "Value Propositions",
    "SWOT Analysis",
    "Business Model Canvas",
    "Business Plan"
]

# -------------------------------
# Predefined prompt templates
# -------------------------------
PROMPTS = {
    "Focus Generation": """You are given the user's story below. Apply the Dilemma Triangle methodology (People, Planet, Prosperity) to extract focus areas.
For each driver, produce only 1 specific focus area and a short rationale (1–2 sentences).
Return only valid JSON and nothing else:
{
  "focuses": [
    {"driver":"People","focus":"...","rationale":"..."},
    {"driver":"Planet","focus":"...","rationale":"..."},
    {"driver":"Prosperity","focus":"...","rationale":"..."}
  ]
}""",

    "Issues Generation": """Given the focus areas (and drivers), list 3–6 issues for each focus area that stem from it.
Return only valid JSON and nothing else:
{
  "issues_by_focus": [
    {"focus":"...","driver":"...","issues":[{"issue":"...","explain":"..."}]}
  ]
}""",

    "Tension Matrix": """Given the issues across focuses, generate a tension matrix describing conflicts or tradeoffs between issues.
Return only valid JSON and nothing else:
{
  "tensions":[
    {"issue_a":"...","issue_b":"...","tension":"...","why":"..."}
  ]
}""",

    "Dilemmas & Ranking": """From the tension matrix, generate dilemmas phrased as tradeoffs.
Each dilemma should include a title, description, affected drivers, and an importance score (1–10).
Return only valid JSON and nothing else:
{
  "dilemmas":[
    {"title":"...","description":"...","drivers":["People","Planet"],"score":8}
  ]
}""",

    "Value Propositions": """For the top dilemmas, propose up to 3 concrete value propositions (solutions) addressing the dilemmas while balancing drivers.
Return only valid JSON and nothing else:
{
  "value_propositions":[
    {"title":"...","explain":"...","dilemmas":["..."],"benefits":["..."]}
  ]
}""",

    "SWOT Analysis": """Perform a SWOT analysis on each provided value proposition.
Return only valid JSON and nothing else:
{
  "swot":[
    {"title":"...","S":["..."],"W":["..."],"O":["..."],"T":["..."]}
  ]
}""",

    "Business Model Canvas": """Generate a Business Model Canvas (9 blocks) for the selected value proposition.
Return only valid JSON and nothing else. Make sure to include all 9 blocks with the exact keys:
- key_partners
- key_activities
- key_resources
- value_propositions
- customer_relationships
- channels
- customer_segments
- revenue_streams
- cost_structure

JSON format example:
{
  "bmc":[
    {
      "value_proposition":"<Title of Value Proposition>",
      "canvas":{
        "key_partners":["..."],
        "key_activities":["..."],
        "key_resources":["..."],
        "value_propositions":["..."],
        "customer_relationships":["..."],
        "channels":["..."],
        "customer_segments":["..."],
        "revenue_streams":["..."],
        "cost_structure":["..."]
      }
    }
  ]
}""",

    "Business Plan": """You are an expert business strategist.

Using all the information provided below — including the original story, SWOT analysis, and the Business Model Canvas — create a clear and structured Business Plan (around 2–3 pages) that includes:

1. Executive Summary
2. Market Opportunity
3. Business Model (connected to the BMC)
4. Product or Service Description
5. Marketing and Customer Strategy
6. Operations Plan
7. Financial and Sustainability Outlook
8. Key Risks and Mitigation Strategies
9. Conclusion and Next Steps

Be concise yet insightful. Use bullet points or short paragraphs where suitable.
Return the business plan in plain text (no JSON or markdown fences)."""
}

# -------------------------------
# Initialize session state
# -------------------------------
if "step_index" not in st.session_state:
    st.session_state.step_index = 0
if "conversation" not in st.session_state:
    st.session_state.conversation = []
if "story" not in st.session_state:
    st.session_state.story = ""
if "completed" not in st.session_state:
    st.session_state.completed = False
if "selected_value_prop" not in st.session_state:
    st.session_state.selected_value_prop = None

# -------------------------------
# Current step
# -------------------------------
current_step = STEPS[st.session_state.step_index]
st.title("🌍 Dilemma Triangle → Business Model Canvas")
# Show step title and selected value proposition if available
if current_step == "Business Plan" and "selected_value_prop" in st.session_state and st.session_state.selected_value_prop:
    vp_title = st.session_state.selected_value_prop.get("title", "")
    st.header(f"Step {st.session_state.step_index + 1}: {current_step} – {vp_title}")
else:
    st.header(f"Step {st.session_state.step_index + 1}: {current_step}")

# -------------------------------
# Step 1: Story input
# -------------------------------
if current_step == "Story Input":
    user_story = st.text_area("✏️ Please provide the full story or context:", value=st.session_state.story, height=200)
    if st.button("Submit Story"):
        if user_story.strip():
            st.session_state.story = user_story.strip()
            st.session_state.conversation.append({
                "step": "Story Input",
                "prompt": user_story.strip(),
                "response": "✅ Story saved successfully.",
                "feedback": ""
            })
            st.session_state.step_index += 1
            st.success("Story submitted. Proceeding to Focus Generation.")
            st.rerun()
        else:
            st.warning("Please enter the story before continuing.")

# -------------------------------
# Step 2–8: LLM-driven steps
# -------------------------------
else:
    if len(st.session_state.conversation) <= st.session_state.step_index:
        prev_outputs = "\n\n".join([f"### Step: {c['step']}\n{c['response']}" for c in st.session_state.conversation])
        base_prompt = PROMPTS.get(current_step, "")
        story_context = st.session_state.story

        # Use selected SWOT if generating Business Model Canvas
        if current_step == "Business Model Canvas" and st.session_state.selected_value_prop:
            selected_swot = st.session_state.selected_value_prop
            prev_outputs = f"### Selected SWOT\n{selected_swot}"
            final_prompt = f"{base_prompt}\n\nContext:\n{story_context}\n\nSelected SWOT Value Proposition:\n{selected_swot}"
        else:
            final_prompt = f"{base_prompt}\n\nContext:\n{story_context}\n\nPrevious Outputs:\n{prev_outputs}"

        with st.spinner(f"Generating {current_step}..."):
            response = model.generate_content(final_prompt)
            text_response = response.text if hasattr(response, "text") else "Error: No valid response."

        st.session_state.conversation.append({
            "step": current_step,
            "prompt": final_prompt,
            "response": text_response,
            "feedback": ""
        })
        st.success(f"✅ {current_step} generated successfully.")
        st.rerun()

    # Display conversation history
    for idx, item in enumerate(st.session_state.conversation):
        st.markdown(f"### {idx + 1}. {item['step']}")
        if item['step'] == "Story Input":
            st.info(f"📖 Story: {item['prompt']}")
        else:
            st.markdown(f"**🤖 LLM Output:**")
            with st.expander("View Output"):
                st.write(item['response'])

        if idx == st.session_state.step_index:
            # Skip feedback/refine/approve for Business Plan step
            if item["step"] == "Business Plan":
                continue
            feedback_key = f"feedback_{idx}"
            if feedback_key not in st.session_state:
                st.session_state[feedback_key] = item.get("feedback", "")

            feedback_text = st.text_input(
                f"Provide feedback for {item['step']}",
                value=st.session_state[feedback_key],
                key=feedback_key + "_input"
            )

            col1, col2 = st.columns(2)
            with col1:
                if st.button(f"🔄 Refine {item['step']}", key=f"refine_{idx}"):

                    if feedback_text.strip():
                        with st.spinner("Refining response..."):
                            refine_prompt = (
                                f"Refine the following output based on this feedback.\n\n"
                                f"Feedback:\n{feedback_text}\n\nOriginal Output:\n{item['response']}"
                            )
                            refined = model.generate_content(refine_prompt)
                            refined_text = refined.text if hasattr(refined, "text") else "Error: No refined response."
                            st.session_state.conversation[idx]["response"] = refined_text
                            st.session_state.conversation[idx]["feedback"] = feedback_text
                        st.success("✅ Response refined successfully.")
                        st.rerun()
                    else:
                        st.warning("Please enter feedback before refining.")

            with col2:
                if st.button(f"✅ Approve {item['step']}", key=f"approve_{idx}"):

                    if st.session_state.step_index < len(STEPS) - 1:
                        st.session_state.step_index += 1
                        st.success(f"Step {idx + 1} approved. Moving to next step: {STEPS[st.session_state.step_index]}")
                        st.rerun()
                    else:
                        st.session_state.completed = True
                        st.success("🎉 All steps completed!")
                        st.rerun()

        else:
            st.caption("✅ Step completed")

# -------------------------------
# SWOT Analysis Visualization + Selection
# -------------------------------
if current_step == "SWOT Analysis" and len(st.session_state.conversation) > 0:
    st.markdown("---")
    st.subheader("🧠 SWOT Analysis Dashboard")

    import json, re

    def listify(value):
        """Convert string or list into bullet list."""
        if isinstance(value, list):
            return value
        elif isinstance(value, str):
            return [v.strip("-• ") for v in re.split(r"[\n,;]", value) if v.strip()]
        return []

    try:
        last_output = st.session_state.conversation[-1]["response"]
        match = re.search(r"(\{(?:.|\n)*\})", last_output)
        json_str = match.group(1) if match else None

        if not json_str:
            st.warning("⚠️ No JSON object found in SWOT output.")
        else:
            data = json.loads(json_str)
            if "swot" in data and isinstance(data["swot"], list):
            
                # ---- Allow user to select one SWOT to continue ----
                options = [entry.get("title", f"Option {i+1}") for i, entry in enumerate(data["swot"])]
                selected_title = st.selectbox("Select the best value proposition to continue:", options)

                st.session_state.selected_value_prop = next(
                    (entry for entry in data["swot"] if entry.get("title") == selected_title),
                    None
                )
                st.success(f"Selected value proposition: {selected_title}")

                for entry in data["swot"]:
                    st.markdown(f"## 🌿 {entry.get('title', 'Untitled Initiative')}")

                    col1, col2 = st.columns(2)
                    with col1:
                        # Strengths
                        st.markdown(
                            """
                            <div style="background-color:#e6ffe6;border-radius:10px;padding:10px 16px;margin-bottom:8px;">
                                <h5>💪 Strengths</h5>
                                <ul style="margin-top:6px;">
                            """ +
                            "".join([f"<li>{s}</li>" for s in listify(entry.get("S"))]) +
                            "</ul></div>",
                            unsafe_allow_html=True,
                        )

                        # Weaknesses
                        st.markdown(
                            """
                            <div style="background-color:#fff0f0;border-radius:10px;padding:10px 16px;margin-bottom:8px;">
                                <h5>⚠️ Weaknesses</h5>
                                <ul style="margin-top:6px;">
                            """ +
                            "".join([f"<li>{w}</li>" for w in listify(entry.get("W"))]) +
                            "</ul></div>",
                            unsafe_allow_html=True,
                        )

                    with col2:
                        # Opportunities
                        st.markdown(
                            """
                            <div style="background-color:#f0f8ff;border-radius:10px;padding:10px 16px;margin-bottom:8px;">
                                <h5>🚀 Opportunities</h5>
                                <ul style="margin-top:6px;">
                            """ +
                            "".join([f"<li>{o}</li>" for o in listify(entry.get("O"))]) +
                            "</ul></div>",
                            unsafe_allow_html=True,
                        )

                        # Threats
                        st.markdown(
                            """
                            <div style="background-color:#fff8e6;border-radius:10px;padding:10px 16px;margin-bottom:8px;">
                                <h5>💣 Threats</h5>
                                <ul style="margin-top:6px;">
                            """ +
                            "".join([f"<li>{t}</li>" for t in listify(entry.get("T"))]) +
                            "</ul></div>",
                            unsafe_allow_html=True,
                        )

                    st.markdown("---")

            else:
                st.info("No valid SWOT data found in output.")
    except json.JSONDecodeError as e:
        st.error(f"❌ Could not parse SWOT JSON: {e}")
    except Exception as e:
        st.error(f"⚠️ Error displaying SWOT Analysis: {e}")

# -------------------------------
# Business Model Canvas Visualization
# -------------------------------
if current_step == "Business Model Canvas" and len(st.session_state.conversation) > 0:
    show_bmc_visualization(st.session_state.conversation[-1]["response"])

# -------------------------------
# Business Plan (view + download only)
# -------------------------------
elif current_step == "Business Plan":
    # The top-level "9. Business Plan" header is already shown — no need to repeat it
    if "selected_value_prop" in st.session_state and st.session_state.selected_value_prop:
        vp_title = st.session_state.selected_value_prop.get("title", "")
        st.markdown(f"### 📄 Business Plan for **{vp_title}**")
    else:
        st.markdown("### 📄 Business Plan")

    # Get the generated Business Plan text from the conversation
    if len(st.session_state.conversation) > 0:
        st.session_state.business_plan = st.session_state.conversation[-1]["response"]

        # Display note and download option
        st.success("✅ Business Plan generated successfully!")
        # Create Word document
        doc = Document()
        doc.add_heading("Business Plan", level=1)
        if "selected_value_prop" in st.session_state and st.session_state.selected_value_prop:
            doc.add_paragraph(f"Value Proposition: {st.session_state.selected_value_prop.get('title','')}")

        doc.add_paragraph(st.session_state.business_plan)

        # Convert to bytes for download
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        st.download_button(
            "⬇️ Download Business Plan (Word)",
            data=buffer,
            file_name="Business_Plan.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    else:
        st.warning("⚠️ No Business Plan found. Please complete the previous steps first.")

